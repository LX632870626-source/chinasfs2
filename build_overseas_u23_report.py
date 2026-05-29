from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUT_DIR = ROOT / "输出"
OUT_DIR.mkdir(exist_ok=True)

IMAGE_SRC = Path(r"C:\Users\xiaoh\.codex\generated_images\019e697e-cb8e-7113-9c66-a6bfce1b4b1b\ig_0a64f5b32966efc0016a16ea0def0081988bed455d1b248548.png")
IMAGE_DST = OUT_DIR / "中国U23球员留洋总结图.png"
DOCX_OUT = OUT_DIR / "2026中国23岁以下足球球员海外效力情况总结.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, bold=False, color="000000", size=9):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if len(text) <= 10 else WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_table_borders(table, color="D9E2EC"):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    style = doc.styles[f"Heading {level}"]
    p.style = style
    run = p.add_run(text)
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    return p


def add_body(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.12
    run = p.add_run(text)
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(10.5)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(10)
    return p


def add_table(doc, headers, rows, widths_cm):
    table = doc.add_table(rows=1, cols=len(headers))
    table.autofit = False
    set_table_borders(table)
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.width = Cm(widths_cm[idx])
        set_cell_shading(cell, "E8F0EA")
        set_cell_text(cell, header, bold=True, color="123322", size=8.5)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].width = Cm(widths_cm[idx])
            set_cell_text(cells[idx], value, size=8.2)
    return table


def main():
    if IMAGE_SRC.exists():
        IMAGE_DST.write_bytes(IMAGE_SRC.read_bytes())

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    styles = doc.styles
    styles["Normal"].font.name = "Microsoft YaHei"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    styles["Normal"].font.size = Pt(10.5)
    for name, size, color in [
        ("Heading 1", 15, "1F4D3A"),
        ("Heading 2", 12.5, "2F6B4F"),
        ("Heading 3", 11.5, "2F6B4F"),
    ]:
        style = styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(2)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("2026年中国23岁以下足球球员海外效力情况总结")
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string("163B2C")

    meta = doc.add_paragraph()
    meta_run = meta.add_run("统计口径：截至2026年5月27日；年龄口径为2003年1月1日以后出生；覆盖海外一线队、预备队/U21、U19/U17梯队及已官宣待加盟。")
    meta_run.font.name = "Microsoft YaHei"
    meta_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    meta_run.font.size = Pt(9)
    meta_run.font.color.rgb = RGBColor.from_string("4B5563")

    if IMAGE_DST.exists():
        doc.add_picture(str(IMAGE_DST), width=Inches(6.5))
        cap = doc.add_paragraph("图：使用 imagegen 生成的留洋路径总结图，已去除俱乐部标识与可读文字。")
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.runs[0].font.size = Pt(8)
        cap.runs[0].font.color.rgb = RGBColor.from_string("68727D")

    add_heading(doc, "一、核心结论", 1)
    for item in [
        "成年化竞争层面仍偏少：2026年中国U23国字号名单中，公开数据可核到的海外球员主要是徐彬和王博豪，两人分别处于英格兰俱乐部体系和荷兰乙级联赛一线队环境；徐彬4月下旬已因腿筋伤势结束巴恩斯利租借并回狼队康复。",
        "低龄梯队数量更可观：2026年AFC U17亚洲杯中国队报名表显示，至少5名2009年龄段球员在西班牙、克罗地亚、塞尔维亚俱乐部梯队注册。",
        "路径正在分化：英格兰路径重视“签约+外租+U21资格”，荷兰路径直接给一线队比赛窗口，西班牙/巴尔干路径更多是青训梯队和青年联赛适应。",
        "最大变量不是“是否出国”，而是海外周期是否能转化为稳定比赛时间、成年队上场和国家队回流价值。",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "二、已在海外或已确定海外路径的球员", 1)
    headers = ["球员", "出生", "位置", "海外俱乐部/梯队", "国家", "状态判断"]
    rows = [
        ["徐彬", "2004-05-02", "后腰", "狼队U21；曾租借巴恩斯利，4月伤退回归狼队", "英格兰", "U23主力级样本；U21出场与康复进展是后续观察点"],
        ["王博豪", "2005-07-18", "后腰/前锋", "FC Den Bosch一线队，陕西联合外租", "荷兰", "一线队联赛样本，数据口径显示已有荷乙出场"],
        ["张家鸣", "2007", "中锋", "伯恩利U21；租借FK Vozdovac U19", "英格兰/塞尔维亚", "英格兰注册+塞尔维亚青年赛练级路径"],
        ["金昱成", "2009-06-15", "后卫", "NK Lokomotiva Zagreb", "克罗地亚", "U17国字号报名表确认海外梯队"],
        ["谢晋", "2009-08-21", "中场", "Real CD Carabanchel", "西班牙", "U17国字号报名表确认海外梯队"],
        ["万项", "2009-04-21", "中场", "Red Star Belgrade", "塞尔维亚", "U17国字号报名表确认海外梯队"],
        ["王修昊", "2009-01-21", "后卫/边路", "Damm CF", "西班牙", "U17国字号报名表确认；俱乐部页可核到球员资料"],
        ["朱宇轩", "2009-07-13", "门将", "Getafe FC", "西班牙", "U17国字号报名表确认海外梯队"],
        ["魏祥鑫", "2008", "中锋", "AJ Auxerre，计划2026-07-01加盟", "法国", "已被多源报道为官宣待到队；截至本报告日尚属待加盟"],
    ]
    add_table(doc, headers, rows, [1.55, 1.65, 1.45, 4.0, 1.55, 5.4])

    add_heading(doc, "三、分层观察", 1)
    add_heading(doc, "1. 成年队/准成年队", 2)
    add_body(doc, "徐彬和王博豪是目前最接近成年队检验的两类样本。徐彬的关键在于伤后康复、狼队U21出场和下一段外租质量；王博豪的关键在于荷乙一线队分钟数是否可持续，并能否形成中场对抗、推进和转换速度方面的真实提升。")
    add_heading(doc, "2. U21/U19过渡层", 2)
    add_body(doc, "张家鸣代表的是另一种务实路线：先进入英格兰俱乐部青年体系，再用巴尔干青年队/低级别赛事补足实战。这个路径的优点是训练体系和比赛环境都更职业化，风险是劳工证、外租质量和回到母队后的队内竞争。")
    add_heading(doc, "3. U17梯队层", 2)
    add_body(doc, "2009年龄段的金昱成、谢晋、万项、王修昊、朱宇轩显示出中国低龄球员进入欧洲梯队的渠道比成年留洋更宽。现阶段评价不宜过度拔高，重点应看三件事：是否长期留队、是否进入更高年龄段主力轮换、是否在18岁前后完成职业合同或成年队梯队晋升。")

    add_heading(doc, "四、机会与风险", 1)
    for item in [
        "机会：海外梯队让球员更早适应高节奏训练、身体对抗、语言文化和职业化竞争。",
        "机会：英格兰、荷兰、法国、塞尔维亚、西班牙、克罗地亚等路径并存，有助于摆脱单一留洋模板。",
        "风险：部分梯队信息公开度低，外界容易把“注册在海外”误读为“已接近一线队”。",
        "风险：U17到U21之间淘汰率极高，真正有价值的节点是成年队分钟数、稳定合同和国家队战术适配。",
        "跟踪指标：联赛/杯赛正式出场分钟、所在年龄段级别、合同年限、是否上调U19/U21/一线队训练、伤病与回国周期。",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "五、主要来源", 1)
    sources = [
        "AFC：2026年AFC U17亚洲杯最终报名表，确认中国U17中金昱成、谢晋、万项、王修昊、朱宇轩的海外俱乐部信息。",
        "Transfermarkt：中国U23 2026名单显示23人中海外球员2名，并列出徐彬为Wolverhampton Wanderers U21、王博豪为FC Den Bosch。",
        "FC Den Bosch官方：王博豪从陕西联合租借加盟FC Den Bosch。",
        "北京日报/懂球帝等：徐彬从狼队租借加盟巴恩斯利，保留狼队U21资格；4月下旬因腿筋伤势结束租借、回狼队康复。",
        "直播吧/新浪等转引俱乐部消息：张家鸣从伯恩利U21租借加盟塞尔维亚FK Vozdovac U19，租期至2026年6月30日。",
        "Transfermarkt专题：魏祥鑫预计2026年夏天前往欧洲，加盟法国欧塞尔。",
        "CF Damm球员页：王修昊在Damm CF的个人资料与注册信息。",
    ]
    for item in sources:
        add_bullet(doc, item)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run("资料整理日期：2026-05-27")
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = RGBColor.from_string("6B7280")

    doc.save(DOCX_OUT)
    print(DOCX_OUT)


if __name__ == "__main__":
    main()
